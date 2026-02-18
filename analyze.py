import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def analyze_django_project(project_path):
    results = []
    file_number = 1
    total_lines = 0
    total_size = 0

    for root, dirs, files in os.walk(project_path):
        for file in files:
            if file.endswith(".py"):
                file_path = os.path.join(root, file)

                # Подсчет строк
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    lines = sum(1 for _ in f)

                # Размер файла в КБ
                size_kb = os.path.getsize(file_path) / 1024
                size_kb = round(size_kb, 2)

                total_lines += lines
                total_size += size_kb

                results.append({
                    "no": file_number,
                    "name": file,
                    "purpose": "Модуль Django",
                    "lines": lines,
                    "size": size_kb
                })

                file_number += 1

    return results, total_lines, round(total_size, 2)


def print_table(data):
    print(f"{'No':<5}{'Наименование модуля':<35}{'Функциональное назначение':<30}{'Строк':<10}{'Размер КБ':<10}")
    print("-" * 100)

    for item in data:
        print(f"{item['no']:<5}"
              f"{item['name']:<35}"
              f"{item['purpose']:<30}"
              f"{item['lines']:<10}"
              f"{item['size']:<10}")


def export_to_word(data, total_lines, total_size, filename="django_project_analysis.docx"):
    document = Document()
    document.add_heading("Анализ Django проекта", level=1)

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = table.rows[0].cells
    headers[0].text = "No"
    headers[1].text = "Наименование модуля"
    headers[2].text = "Функциональное назначение"
    headers[3].text = "Количество строк"
    headers[4].text = "Размер (КБ)"

    for item in data:
        row = table.add_row().cells
        row[0].text = str(item["no"])
        row[1].text = item["name"]
        row[2].text = item["purpose"]
        row[3].text = str(item["lines"])
        row[4].text = str(item["size"])

    document.add_paragraph("")
    document.add_paragraph(f"Общее количество строк кода: {total_lines}")
    document.add_paragraph(f"Общий размер проекта: {total_size} КБ")

    document.save(filename)
    print(f"\nWord файл создан: {filename}")


if __name__ == "__main__":
    project_path = "."  # путь к проекту (где manage.py)
    data, total_lines, total_size = analyze_django_project(project_path)

    print_table(data)
    export_to_word(data, total_lines, total_size)
