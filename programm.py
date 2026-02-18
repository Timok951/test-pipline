import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_project_code_to_word(project_path, output_file="django_full_code.docx"):
    document = Document()
    document.add_heading("Листинг программного кода проекта Django", level=1)

    for root, dirs, files in os.walk(project_path):
        for file in files:
            if file.endswith(".py"):
                file_path = os.path.join(root, file)

                # Заголовок файла
                document.add_heading(f"Файл: {file}", level=2)
                document.add_paragraph(f"Путь: {file_path}")

                # Читаем код
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    code = f.read()

                # Добавляем код как моноширинный текст
                paragraph = document.add_paragraph()
                run = paragraph.add_run(code)
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
                run.font.size = Pt(9)

                document.add_page_break()

    document.save(output_file)
    print(f"Готово! Файл создан: {os.path.abspath(output_file)}")


if __name__ == "__main__":
    project_path = "."  # Запускать из корня проекта (где manage.py)
    export_project_code_to_word(project_path)
