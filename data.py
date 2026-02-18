#!/usr/bin/env python3
"""
Собирает все классы из .py файлов Django-проекта (кроме миграций).
Выводит в консоль + сохраняет в документ Word: classes_in_project.docx
"""

import ast
from pathlib import Path
import sys
from typing import List, Dict

# ────────────────────────────────────────────────
# Блок для Word-документа
# ────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    print("⚠️  Библиотека python-docx не установлена.")
    print("    Установите её командой:  pip install python-docx")
    print("    Документ Word создан НЕ будет, только консольный вывод.\n")

# ────────────────────────────────────────────────
# Функции парсинга (без изменений)
# ────────────────────────────────────────────────

def get_base_names(bases) -> List[str]:
    result = []
    for base in bases:
        if isinstance(base, ast.Name):
            result.append(base.id)
        elif isinstance(base, ast.Attribute):
            result.append(base.attr)
        elif isinstance(base, ast.Subscript):
            if hasattr(base.value, 'id'):
                result.append(base.value.id)
            else:
                result.append("Generic[...]")  
        else:
            result.append("<?>")
    return result


def extract_classes_from_file(filepath: Path) -> List[Dict]:
    classes = []
    try:
        with filepath.open("r", encoding="utf-8") as f:
            content = f.read()
        tree = ast.parse(content, filename=str(filepath))
    except Exception as e:
        print(f"Ошибка парсинга {filepath}: {e}", file=sys.stderr)
        return []

    for node in ast.iter_child_nodes(tree):
        if not isinstance(node, ast.ClassDef):
            continue

        bases = get_base_names(node.bases)
        class_info = {
            "name": node.name,
            "bases": bases,
            "file": str(filepath.relative_to(project_root)),
            "line": node.lineno,
            "docstring": (ast.get_docstring(node) or "").strip(),
        }
        classes.append(class_info)

    return classes


def is_relevant_file(path: Path) -> bool:
    name = path.name
    if name == "__init__.py":
        return False
    if "migrations" in path.parts:
        return False
    return name.endswith(".py")


# ────────────────────────────────────────────────
# Основная логика
# ────────────────────────────────────────────────

if len(sys.argv) < 2:
    print("Использование:")
    print("  python collect_django_classes_to_word.py .")
    print("  python collect_django_classes_to_word.py /path/to/project")
    sys.exit(1)

project_root = Path(sys.argv[1]).resolve()
if not project_root.is_dir():
    print(f"Путь не является директорией: {project_root}")
    sys.exit(1)

print(f"Сканирую проект: {project_root}\n")

all_classes: List[Dict] = []

for py_file in project_root.rglob("*.py"):
    if not is_relevant_file(py_file):
        continue
    classes_in_file = extract_classes_from_file(py_file)
    all_classes.extend(classes_in_file)

all_classes.sort(key=lambda c: (c["file"], c["line"]))

# ────────────────────────────────────────────────
# Вывод в консоль
# ────────────────────────────────────────────────

print(f"Найдено классов: {len(all_classes)}\n")

if not all_classes:
    print("Классы не найдены.")
else:
    current_file = None
    for cls in all_classes:
        if cls["file"] != current_file:
            current_file = cls["file"]
            print(f"\n{current_file}")
            print("─" * 60)

        bases_str = ", ".join(cls["bases"]) if cls["bases"] else "—"
        print(f"  {cls['name']:<40}  (наследует: {bases_str})  стр. {cls['line']:>4}")
        
        if cls["docstring"]:
            first_line = cls["docstring"].split("\n")[0][:110]
            print(f"      # {first_line}...")
        else:
            print("      # (без docstring)")

# ────────────────────────────────────────────────
# Создание Word-документа
# ────────────────────────────────────────────────

# ────────────────────────────────────────────────
# Создание Word-документа с ТАБЛИЦЕЙ
# ────────────────────────────────────────────────
if WORD_AVAILABLE and all_classes:
    doc = Document()
    doc.add_heading(f"Классы Django-проекта — {project_root.name}", level=1)

    p = doc.add_paragraph(f"Сканировано: {project_root}\nВсего классов: {len(all_classes)}\nДата: ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавляем таблицу
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'  # или 'Light Grid Accent 1' и т.д.
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ п/п'
    hdr_cells[1].text = 'Наименование'
    hdr_cells[2].text = 'Описание'

    # Делаем заголовок жирным и центрированным
    for cell in hdr_cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Заполняем строки
    for idx, cls in enumerate(all_classes, start=1):
        row_cells = table.add_row().cells

        row_cells[0].text = str(idx)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Наименование + наследование + файл + строка
        name_run = row_cells[1].paragraphs[0].add_run(f"{cls['name']}")
        name_run.bold = True
        row_cells[1].paragraphs[0].add_run(f"\n(наследует: {', '.join(cls['bases']) if cls['bases'] else '—'})\n")
        row_cells[1].paragraphs[0].add_run(f"файл: {cls['file']}  стр. {cls['line']}")

        # Описание
        if cls["docstring"]:
            desc = cls["docstring"].split("\n")[0].strip()[:250]
            if len(cls["docstring"].split("\n")[0]) > 250:
                desc += "..."
            row_cells[2].text = desc
        else:
            row_cells[2].text = "(без docstring / описания)"

    # Авто-подгонка ширины столбцов (примерно)
    table.columns[0].width = Inches(0.8)
    table.columns[1].width = Inches(3.0)
    table.columns[2].width = Inches(4.0)

    output_path = Path("classes_table_project.docx").resolve()
    doc.save(output_path)
    print(f"\nТаблица сохранена в: {output_path}")